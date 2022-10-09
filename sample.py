from docx import Document
from toc import insert_toc, update_toc

d = Document()
d.add_heading('Title', 1)
d.add_heading('Subtitle', 2)
insert_toc(d)
f = '/tmp/demo.docx'
d.save(f)

# update toc
update_toc(f)
